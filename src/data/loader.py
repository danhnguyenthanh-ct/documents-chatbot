"""
Document Loading Module
Supports multiple document formats with metadata extraction and error handling.
"""

import logging
from typing import List, Dict, Any, Optional, Union, BinaryIO
from pathlib import Path
import mimetypes
from datetime import datetime
import hashlib

import PyPDF2
from docx import Document
import chardet
from bs4 import BeautifulSoup
import markdown
import re

logger = logging.getLogger(__name__)


class DocumentLoadError(Exception):
    """Custom exception for document loading operations"""
    pass


class DocumentMetadata:
    """Document metadata container"""
    
    def __init__(
        self,
        file_path: str,
        file_type: str,
        file_size: int,
        created_at: datetime,
        content_hash: str,
        page_count: Optional[int] = None,
        title: Optional[str] = None,
        author: Optional[str] = None,
        **kwargs
    ):
        self.file_path = file_path
        self.file_type = file_type
        self.file_size = file_size
        self.created_at = created_at
        self.content_hash = content_hash
        self.page_count = page_count
        self.title = title
        self.author = author
        self.custom_metadata = kwargs
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert metadata to dictionary"""
        return {
            "file_path": self.file_path,
            "file_type": self.file_type,
            "file_size": self.file_size,
            "created_at": self.created_at.isoformat(),
            "content_hash": self.content_hash,
            "page_count": self.page_count,
            "title": self.title,
            "author": self.author,
            **self.custom_metadata
        }


class LoadedDocument:
    """Container for loaded document content and metadata"""
    
    def __init__(self, content: str, metadata: DocumentMetadata):
        self.content = content
        self.metadata = metadata
    
    def __len__(self) -> int:
        return len(self.content)


class DocumentLoader:
    """
    Multi-format document loader with metadata extraction
    """
    
    SUPPORTED_FORMATS = {
        ".pdf": "application/pdf",
        ".txt": "text/plain",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".html": "text/html",
        ".htm": "text/html",
        ".md": "text/markdown",
        ".markdown": "text/markdown"
    }
    
    def __init__(self, max_file_size: int = 50 * 1024 * 1024):  # 50MB default
        """
        Initialize document loader
        
        Args:
            max_file_size: Maximum file size in bytes
        """
        self.max_file_size = max_file_size
        self.stats = {
            "files_processed": 0,
            "total_size_processed": 0,
            "successful_loads": 0,
            "failed_loads": 0
        }
    
    def is_supported_format(self, file_path: Union[str, Path]) -> bool:
        """Check if file format is supported"""
        file_path = Path(file_path)
        return file_path.suffix.lower() in self.SUPPORTED_FORMATS
    
    def _detect_encoding(self, file_path: Path) -> str:
        """Detect text file encoding"""
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)  # Read first 10KB for detection
                result = chardet.detect(raw_data)
                return result.get('encoding', 'utf-8') or 'utf-8'
        except Exception:
            return 'utf-8'
    
    def _calculate_content_hash(self, content: str) -> str:
        """Calculate hash of document content"""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()
    
    def _load_pdf(self, file_path: Path) -> LoadedDocument:
        """Load PDF document"""
        try:
            text_content = []
            page_count = 0
            title = None
            author = None
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                # Extract metadata
                if pdf_reader.metadata:
                    title = pdf_reader.metadata.get('/Title')
                    author = pdf_reader.metadata.get('/Author')
                
                # Extract text from all pages
                for page in pdf_reader.pages:
                    text_content.append(page.extract_text())
            
            content = '\n'.join(text_content)
            
            if not content.strip():
                raise DocumentLoadError("PDF contains no extractable text")
            
            # Create metadata
            file_stats = file_path.stat()
            metadata = DocumentMetadata(
                file_path=str(file_path),
                file_type="pdf",
                file_size=file_stats.st_size,
                created_at=datetime.fromtimestamp(file_stats.st_ctime),
                content_hash=self._calculate_content_hash(content),
                page_count=page_count,
                title=title,
                author=author
            )
            
            return LoadedDocument(content, metadata)
            
        except Exception as e:
            logger.error(f"Failed to load PDF {file_path}: {e}")
            raise DocumentLoadError(f"PDF loading failed: {e}")
    
    def _load_docx(self, file_path: Path) -> LoadedDocument:
        """Load DOCX document"""
        try:
            doc = Document(file_path)
            
            # Extract text content
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            content = '\n'.join(paragraphs)
            
            if not content.strip():
                raise DocumentLoadError("DOCX contains no text content")
            
            # Extract metadata
            title = doc.core_properties.title
            author = doc.core_properties.author
            
            # Create metadata
            file_stats = file_path.stat()
            metadata = DocumentMetadata(
                file_path=str(file_path),
                file_type="docx",
                file_size=file_stats.st_size,
                created_at=datetime.fromtimestamp(file_stats.st_ctime),
                content_hash=self._calculate_content_hash(content),
                page_count=len(doc.paragraphs),
                title=title,
                author=author
            )
            
            return LoadedDocument(content, metadata)
            
        except Exception as e:
            logger.error(f"Failed to load DOCX {file_path}: {e}")
            raise DocumentLoadError(f"DOCX loading failed: {e}")
    
    def _load_text(self, file_path: Path) -> LoadedDocument:
        """Load text document"""
        try:
            encoding = self._detect_encoding(file_path)
            
            with open(file_path, 'r', encoding=encoding) as file:
                content = file.read()
            
            if not content.strip():
                raise DocumentLoadError("Text file is empty")
            
            # Create metadata
            file_stats = file_path.stat()
            metadata = DocumentMetadata(
                file_path=str(file_path),
                file_type="text",
                file_size=file_stats.st_size,
                created_at=datetime.fromtimestamp(file_stats.st_ctime),
                content_hash=self._calculate_content_hash(content),
                encoding=encoding
            )
            
            return LoadedDocument(content, metadata)
            
        except Exception as e:
            logger.error(f"Failed to load text file {file_path}: {e}")
            raise DocumentLoadError(f"Text loading failed: {e}")
    
    def _load_html(self, file_path: Path) -> LoadedDocument:
        """Load HTML document"""
        try:
            encoding = self._detect_encoding(file_path)
            
            with open(file_path, 'r', encoding=encoding) as file:
                html_content = file.read()
            
            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'lxml')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text content
            content = soup.get_text()
            content = '\n'.join(line.strip() for line in content.splitlines() if line.strip())
            
            if not content.strip():
                raise DocumentLoadError("HTML contains no text content")
            
            # Extract title
            title = soup.title.string if soup.title else None
            
            # Create metadata
            file_stats = file_path.stat()
            metadata = DocumentMetadata(
                file_path=str(file_path),
                file_type="html",
                file_size=file_stats.st_size,
                created_at=datetime.fromtimestamp(file_stats.st_ctime),
                content_hash=self._calculate_content_hash(content),
                title=title,
                encoding=encoding
            )
            
            return LoadedDocument(content, metadata)
            
        except Exception as e:
            logger.error(f"Failed to load HTML {file_path}: {e}")
            raise DocumentLoadError(f"HTML loading failed: {e}")
    
    def _load_markdown(self, file_path: Path) -> LoadedDocument:
        """Load Markdown document"""
        try:
            encoding = self._detect_encoding(file_path)
            
            with open(file_path, 'r', encoding=encoding) as file:
                markdown_content = file.read()
            
            if not markdown_content.strip():
                raise DocumentLoadError("Markdown file is empty")
            
            # Parse markdown and extract metadata
            md = markdown.Markdown(extensions=['meta', 'toc'])
            html_content = md.convert(markdown_content)
            
            # Extract plain text from HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            content = soup.get_text()
            content = '\n'.join(line.strip() for line in content.splitlines() if line.strip())
            
            if not content.strip():
                raise DocumentLoadError("Markdown contains no text content")
            
            # Extract markdown-specific metadata
            title = None
            author = None
            headers = []
            links = []
            
            # Extract title from metadata or first H1
            if hasattr(md, 'Meta') and md.Meta:
                title = md.Meta.get('title', [None])[0] if 'title' in md.Meta else None
                author = md.Meta.get('author', [None])[0] if 'author' in md.Meta else None
            
            # If no title in metadata, extract from first header
            if not title:
                h1_match = re.search(r'^#\s+(.+)$', markdown_content, re.MULTILINE)
                if h1_match:
                    title = h1_match.group(1).strip()
            
            # Extract headers (for structure analysis)
            header_matches = re.findall(r'^(#{1,6})\s+(.+)$', markdown_content, re.MULTILINE)
            headers = [{'level': len(match[0]), 'text': match[1].strip()} for match in header_matches]
            
            # Extract links
            link_matches = re.findall(r'\[([^\]]+)\]\(([^)]+)\)', markdown_content)
            links = [{'text': match[0], 'url': match[1]} for match in link_matches]
            
            # Count various markdown elements
            code_blocks = len(re.findall(r'```[^`]*```', markdown_content, re.DOTALL))
            inline_code = len(re.findall(r'`[^`]+`', markdown_content))
            images = len(re.findall(r'!\[([^\]]*)\]\(([^)]+)\)', markdown_content))
            tables = len(re.findall(r'\|.*\|', markdown_content))
            
            # Create metadata
            file_stats = file_path.stat()
            metadata = DocumentMetadata(
                file_path=str(file_path),
                file_type="markdown",
                file_size=file_stats.st_size,
                created_at=datetime.fromtimestamp(file_stats.st_ctime),
                content_hash=self._calculate_content_hash(content),
                title=title,
                author=author,
                encoding=encoding,
                # Markdown-specific metadata
                header_count=len(headers),
                headers=headers[:10],  # Limit to first 10 headers
                link_count=len(links),
                links=links[:20],  # Limit to first 20 links
                code_block_count=code_blocks,
                inline_code_count=inline_code,
                image_count=images,
                table_count=tables,
                has_toc=bool(re.search(r'\[TOC\]', markdown_content, re.IGNORECASE)),
                raw_markdown_length=len(markdown_content)
            )
            
            return LoadedDocument(content, metadata)
            
        except Exception as e:
            logger.error(f"Failed to load Markdown {file_path}: {e}")
            raise DocumentLoadError(f"Markdown loading failed: {e}")
    
    def load_document(self, file_path: Union[str, Path]) -> LoadedDocument:
        """
        Load a single document
        
        Args:
            file_path: Path to the document file
            
        Returns:
            LoadedDocument: Document content and metadata
            
        Raises:
            DocumentLoadError: If loading fails
        """
        file_path = Path(file_path)
        
        # Validate file exists
        if not file_path.exists():
            raise DocumentLoadError(f"File not found: {file_path}")
        
        # Validate file size
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            raise DocumentLoadError(f"File too large: {file_size} bytes (max: {self.max_file_size})")
        
        # Validate format
        if not self.is_supported_format(file_path):
            raise DocumentLoadError(f"Unsupported format: {file_path.suffix}")
        
        self.stats["files_processed"] += 1
        self.stats["total_size_processed"] += file_size
        
        try:
            # Load based on file type
            suffix = file_path.suffix.lower()
            
            if suffix == ".pdf":
                document = self._load_pdf(file_path)
            elif suffix == ".docx":
                document = self._load_docx(file_path)
            elif suffix in [".txt"]:
                document = self._load_text(file_path)
            elif suffix in [".html", ".htm"]:
                document = self._load_html(file_path)
            elif suffix in [".md", ".markdown"]:
                document = self._load_markdown(file_path)
            else:
                raise DocumentLoadError(f"Handler not implemented for {suffix}")
            
            self.stats["successful_loads"] += 1
            logger.info(f"Successfully loaded document: {file_path}")
            return document
            
        except Exception as e:
            self.stats["failed_loads"] += 1
            raise
    
    def load_documents(self, file_paths: List[Union[str, Path]]) -> List[LoadedDocument]:
        """
        Load multiple documents
        
        Args:
            file_paths: List of file paths
            
        Returns:
            List[LoadedDocument]: Successfully loaded documents
        """
        documents = []
        
        for file_path in file_paths:
            try:
                document = self.load_document(file_path)
                documents.append(document)
            except DocumentLoadError as e:
                logger.warning(f"Skipped document {file_path}: {e}")
                continue
        
        logger.info(f"Loaded {len(documents)} out of {len(file_paths)} documents")
        return documents
    
    def get_stats(self) -> Dict[str, Any]:
        """Get loader statistics"""
        return self.stats.copy()


def create_document_loader(max_file_size: int = 50 * 1024 * 1024) -> DocumentLoader:
    """
    Factory function to create document loader
    
    Args:
        max_file_size: Maximum file size in bytes
        
    Returns:
        DocumentLoader: Configured document loader
    """
    return DocumentLoader(max_file_size=max_file_size) 